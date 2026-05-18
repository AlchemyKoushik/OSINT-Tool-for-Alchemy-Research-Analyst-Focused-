from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path("OSINT_Tool_Alpha_Trends_Technical_Architecture.docx")

TITLE_COLOR = RGBColor(0x0F, 0x3D, 0x56)
ACCENT_COLOR = RGBColor(0x1F, 0x6F, 0x8B)
MUTED_COLOR = RGBColor(0x5F, 0x6B, 0x73)
LIGHT_FILL = "EAF4F8"
ACCENT_FILL = "D8EEF5"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table, color="B7C6CE", size="8"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def add_run(paragraph, text, *, bold=False, color=None, size=None, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)
    return run


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = TITLE_COLOR if level == 1 else ACCENT_COLOR
    run.font.size = Pt(18 if level == 1 else 14)
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    run.font.size = Pt(11)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.1
        run = paragraph.add_run(item)
        run.font.size = Pt(11)


def build_cover(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(60)
    title.paragraph_format.space_after = Pt(12)
    add_run(title, "OSINT Tool Alpha Build", bold=True, color=TITLE_COLOR, size=24)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)
    add_run(subtitle, "Technical Architecture", bold=True, color=ACCENT_COLOR, size=20)
    add_run(subtitle, " - Trends Workflow", bold=True, color=ACCENT_COLOR, size=20)

    strap = doc.add_paragraph()
    strap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    strap.paragraph_format.space_after = Pt(26)
    add_run(
        strap,
        "Meeting-ready architecture note focused on the current Trends-only research path.",
        color=MUTED_COLOR,
        size=11,
        italic=True,
    )

    summary_box = doc.add_table(rows=1, cols=1)
    summary_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary_box.autofit = True
    cell = summary_box.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    add_run(p, "Purpose: ", bold=True, color=TITLE_COLOR, size=11)
    add_run(
        p,
        "Explain how the alpha system generates source-backed market trends through a staged OSINT pipeline.",
        size=11,
    )
    set_table_borders(summary_box, color="C9D9E1", size="10")

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(meta, "Prepared for Alpha Demo Discussion", bold=True, color=MUTED_COLOR, size=10)

    doc.add_section(WD_SECTION.NEW_PAGE)


def build_overview_table(doc):
    add_heading(doc, "1. Executive Overview", level=1)
    add_body(
        doc,
        "This alpha build is designed as a Trends-first OSINT research workflow. The system accepts a topic and optional geography, generates trend-oriented search queries, retrieves and filters sources, collects text from web pages and PDFs, cleans the evidence, and produces structured trend outputs with traceable sources.",
    )

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    headers = ("Layer", "Responsibility")
    for idx, value in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, ACCENT_FILL)
        p = cell.paragraphs[0]
        add_run(p, value, bold=True, color=TITLE_COLOR, size=11)

    rows = [
        ("Frontend UI", "Captures analyst input, sends research requests, and renders trends results."),
        ("API Layer", "Validates requests, coordinates sessions, and exposes research and follow-up endpoints."),
        ("Research Orchestrator", "Runs the end-to-end research pipeline in one bounded execution path."),
        ("Retrieval Layer", "Generates queries, searches the web, filters results, and prioritizes quality sources."),
        ("Evidence Layer", "Collects artifacts, cleans text, scores source quality, and prepares analysis context."),
        ("Analysis Layer", "Produces structured trend insights and attaches source references."),
    ]

    for layer, responsibility in rows:
        cells = table.add_row().cells
        cells[0].text = layer
        cells[1].text = responsibility
        for index, cell in enumerate(cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if index == 0:
                set_cell_shading(cell, LIGHT_FILL)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10.5)

    set_table_borders(table)


def build_flow(doc):
    add_heading(doc, "2. Trends Workflow", level=1)
    add_body(doc, "The current trends pipeline can be described in the following architecture sequence:")

    flow_box = doc.add_table(rows=1, cols=1)
    flow_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = flow_box.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Topic + Geography Input", bold=True, color=TITLE_COLOR, size=11)
    add_run(p, "\n-> Query Generation for Trends", size=11)
    add_run(p, "\n-> Multi-Query Search and Ranking", size=11)
    add_run(p, "\n-> Web / PDF Artifact Collection", size=11)
    add_run(p, "\n-> Evidence Cleaning and Sentence Selection", size=11)
    add_run(p, "\n-> Structured Trend Analysis", size=11)
    add_run(p, "\n-> Source Attachment and Example Enrichment", size=11)
    add_run(p, "\n-> Frontend Trend Cards and Export", size=11)
    set_table_borders(flow_box, color="C9D9E1", size="10")

    add_heading(doc, "3. Stage-by-Stage Breakdown", level=1)
    add_bullets(
        doc,
        [
            "Input Layer: The analyst selects the topic, the section as Trends, and an optional global, regional, or country scope.",
            "Query Generation: The backend creates exactly ten trend-oriented search queries designed to capture what is changing in the market.",
            "Search and Ranking: The system fans out across multiple queries, filters weak domains, blocks reference-only sources such as Wikipedia-family domains, and ranks results by authority, freshness, and geography relevance.",
            "Artifact Collection: The pipeline captures usable text from web pages and PDFs and stores research artifacts for reuse in later analysis or follow-up flows.",
            "Evidence Cleaning: Raw source text is cleaned, noise is removed, and high-signal evidence sentences are selected to form a bounded analysis pack.",
            "Trend Analysis: The cleaned evidence is sent to the analysis layer, which generates structured trend insights rather than plain conversational summaries.",
            "Source Linking: Each returned trend item is normalized and paired with source references so the frontend can render traceable results.",
            "Example Enrichment: A second pass can research concrete examples for top trend items without forcing examples from the same initial evidence bundle.",
        ],
    )


def build_component_details(doc):
    add_heading(doc, "4. Key Components in the Current Alpha Build", level=1)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ("Component", "Primary Role", "Representative Implementation")
    for idx, value in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, ACCENT_FILL)
        add_run(cell.paragraphs[0], value, bold=True, color=TITLE_COLOR, size=11)

    rows = [
        ("Frontend", "Input capture, request orchestration, result rendering", "Static UI with research and follow-up flows"),
        ("FastAPI Backend", "Endpoint routing and orchestration entrypoint", "/api/research and related routes"),
        ("Query Generator", "Builds ten trend-focused search queries", "Structured and fallback query generation"),
        ("Search Service", "Runs DDG search, filtering, scoring, and URL dedupe", "Authority, freshness, and location-aware ranking"),
        ("Scraper Service", "Collects web and PDF artifacts for evidence capture", "Artifact retrieval and storage pipeline"),
        ("Content Processor", "Cleans text and extracts strong evidence sentences", "Bounded context construction for analysis"),
        ("OpenAI Analysis Layer", "Produces structured trends output", "Validated, schema-based trend generation"),
        ("Session and Storage Layer", "Maintains session state and persists research assets", "Redis plus Cloudflare R2"),
    ]

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if idx == 0:
                set_cell_shading(cells[idx], LIGHT_FILL)
            for p in cells[idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    set_table_borders(table)


def build_alpha_notes(doc):
    add_heading(doc, "5. Alpha Build Positioning", level=1)
    add_body(
        doc,
        "For stakeholder discussion, this build should be presented as a production-shaped alpha. The architecture is modular and traceability-focused, but some dependencies and operational controls are still early-stage.",
    )
    add_bullets(
        doc,
        [
            "Strength: The system is structured as a staged research workflow rather than a one-shot summarizer.",
            "Strength: Retrieval, cleaning, analysis, and follow-up decisioning are separated into distinct layers.",
            "Strength: Trend outputs are designed to be source-backed and structured for repeatable rendering.",
            "Alpha Caveat: Reliability still depends on external services such as OpenAI, search providers, and artifact storage.",
            "Alpha Caveat: Fallback behavior exists for several failures, but enterprise-scale hardening is still a future step.",
            "Alpha Caveat: The current story is strongest when positioned as an analyst-assist research engine focused on trend discovery.",
        ],
    )

    add_heading(doc, "6. Suggested Meeting Talk Track", level=1)
    quote_box = doc.add_table(rows=1, cols=1)
    quote_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = quote_box.cell(0, 0)
    set_cell_shading(cell, ACCENT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    add_run(
        p,
        'This alpha build supports a Trends-first OSINT workflow. A user submits a topic, the backend generates trend-specific research queries, retrieves and filters sources, captures web and PDF artifacts, cleans the evidence, and then uses structured AI analysis to return source-backed market trends. The value of the system is that it behaves like a bounded research pipeline rather than a generic chat summary layer.',
        size=11,
    )
    set_table_borders(quote_box, color="9FC6D3", size="10")


def add_header_footer(doc):
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if not header.text:
            add_run(header, "OSINT Tool Alpha Build | Trends Technical Architecture", color=MUTED_COLOR, size=9)

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if not footer.text:
            add_run(footer, "Alpha Architecture Note", color=MUTED_COLOR, size=9)


def main():
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(11)

    build_cover(doc)
    build_overview_table(doc)
    build_flow(doc)
    build_component_details(doc)
    build_alpha_notes(doc)
    add_header_footer(doc)

    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH.resolve())


if __name__ == "__main__":
    main()
